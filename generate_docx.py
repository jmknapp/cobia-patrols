#!/usr/bin/env python3
"""
Convert cobia_story.md to a Word document.
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def parse_markdown(md_text):
    """Parse markdown text into structured elements."""
    lines = md_text.split('\n')
    elements = []
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Heading 1
        if line.startswith('# '):
            elements.append(('h1', line[2:].strip()))
            i += 1
            continue
        
        # Heading 2
        if line.startswith('## '):
            elements.append(('h2', line[3:].strip()))
            i += 1
            continue
        
        # Heading 3
        if line.startswith('### '):
            elements.append(('h3', line[4:].strip()))
            i += 1
            continue
        
        # Horizontal rule
        if line.strip() == '---':
            elements.append(('hr', ''))
            i += 1
            continue
        
        # Blockquote (collect all consecutive blockquote lines)
        if line.startswith('> '):
            quote_lines = []
            while i < len(lines) and lines[i].startswith('> '):
                quote_lines.append(lines[i][2:])
                i += 1
            elements.append(('blockquote', ' '.join(quote_lines)))
            continue
        
        # Bullet list item
        if line.startswith('- '):
            elements.append(('bullet', line[2:].strip()))
            i += 1
            continue
        
        # Indented bullet (sub-item)
        if line.startswith('  - '):
            elements.append(('bullet2', line[4:].strip()))
            i += 1
            continue
        
        # Regular paragraph (collect lines until empty line or special marker)
        para_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and not lines[i].startswith('- ') and not lines[i].startswith('> ') and lines[i].strip() != '---':
            para_lines.append(lines[i])
            i += 1
        if para_lines:
            elements.append(('para', ' '.join(para_lines)))
        continue
    
    return elements

def add_formatted_text(paragraph, text):
    """Add text with basic formatting (bold, italic) to a paragraph."""
    # Pattern to match **bold**, *italic*, and [link](url)
    pattern = r'(\*\*.*?\*\*|\*.*?\*|\[.*?\]\(.*?\))'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            # Bold
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Italic
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('[') and '](' in part:
            # Link - just show the text
            match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                link_text = match.group(1)
                run = paragraph.add_run(link_text)
                run.underline = True
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)

def create_docx(md_file, output_file):
    """Create a Word document from a markdown file."""
    
    # Read markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        md_text = f.read()
    
    # Parse markdown
    elements = parse_markdown(md_text)
    
    # Create document
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Modify Normal style
    style = styles['Normal']
    style.font.name = 'Georgia'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15
    
    # Add elements
    for elem_type, content in elements:
        if elem_type == 'h1':
            p = doc.add_heading(content, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        elif elem_type == 'h2':
            doc.add_heading(content, level=2)
        
        elif elem_type == 'h3':
            doc.add_heading(content, level=3)
        
        elif elem_type == 'hr':
            # Add a subtle separator
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('— ✦ —')
            run.font.size = Pt(12)
        
        elif elem_type == 'blockquote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            add_formatted_text(p, content)
            for run in p.runs:
                run.italic = True
        
        elif elem_type == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, content)
        
        elif elem_type == 'bullet2':
            p = doc.add_paragraph(style='List Bullet 2')
            add_formatted_text(p, content)
        
        elif elem_type == 'para':
            p = doc.add_paragraph()
            add_formatted_text(p, content)
    
    # Save document
    doc.save(output_file)
    print(f"Word document created: {output_file}")

if __name__ == '__main__':
    create_docx('cobia_story.md', 'cobia_story.docx')



945:




