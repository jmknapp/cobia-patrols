#!/usr/bin/env python3
"""
OCR a patrol report page image and create a Word document 
that preserves the original text layout and formatting.
"""

import sys
import pytesseract
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def get_ocr_with_layout(image_path):
    """Get OCR text with layout preservation using Tesseract's built-in layout analysis."""
    img = Image.open(image_path)
    
    # Use psm 6 for uniform block of text, or psm 4 for single column
    # psm 6 = Assume a single uniform block of text
    # psm 4 = Assume a single column of text of variable sizes
    # psm 3 = Fully automatic page segmentation (default)
    custom_config = r'--psm 6'
    
    # Get data at the line level using block/par/line structure
    data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT, config=custom_config)
    
    return data, img.size

def reconstruct_lines(data, img_width):
    """Reconstruct lines from OCR data, properly grouping words."""
    
    # Group by block_num, par_num, line_num
    lines_dict = {}
    
    n = len(data['text'])
    for i in range(n):
        text = data['text'][i].strip()
        if not text:
            continue
        
        block = data['block_num'][i]
        par = data['par_num'][i]
        line = data['line_num'][i]
        
        key = (block, par, line)
        
        if key not in lines_dict:
            lines_dict[key] = {
                'words': [],
                'top': data['top'][i],
                'left': float('inf'),
                'right': 0
            }
        
        lines_dict[key]['words'].append({
            'text': text,
            'left': data['left'][i],
            'width': data['width'][i],
            'conf': data['conf'][i]
        })
        lines_dict[key]['left'] = min(lines_dict[key]['left'], data['left'][i])
        lines_dict[key]['right'] = max(lines_dict[key]['right'], data['left'][i] + data['width'][i])
    
    # Convert to list and sort by vertical position
    lines = []
    for key in sorted(lines_dict.keys()):
        line_data = lines_dict[key]
        # Sort words by horizontal position
        words = sorted(line_data['words'], key=lambda w: w['left'])
        
        # Reconstruct line text with spacing
        line_text = ""
        prev_right = 0
        for i, word in enumerate(words):
            if i == 0:
                line_text = word['text']
                prev_right = word['left'] + word['width']
            else:
                gap = word['left'] - prev_right
                if gap > 50:  # Large gap - likely tab
                    line_text += "\t" + word['text']
                elif gap > 20:  # Medium gap - multiple spaces
                    spaces = max(2, gap // 10)
                    line_text += " " * spaces + word['text']
                else:
                    line_text += " " + word['text']
                prev_right = word['left'] + word['width']
        
        lines.append({
            'text': line_text,
            'left': line_data['left'],
            'right': line_data['right'],
            'top': line_data['top'],
            'key': key
        })
    
    # Sort by top position
    lines.sort(key=lambda x: x['top'])
    
    return lines

def analyze_lines(lines, img_width):
    """Analyze lines to determine formatting."""
    if not lines:
        return []
    
    # Find typical left margin
    lefts = [l['left'] for l in lines]
    left_margin = min(lefts)
    
    # Typical indentation unit (roughly 0.5 inch in pixels at ~100 DPI)
    indent_unit = 40
    
    analyzed = []
    for line in lines:
        text = line['text']
        
        # Calculate indent level
        indent_pixels = line['left'] - left_margin
        indent_level = int(indent_pixels / indent_unit)
        indent_level = min(indent_level, 6)  # Cap at 6 levels
        
        # Check if centered
        line_center = (line['left'] + line['right']) / 2
        page_center = img_width / 2
        is_centered = abs(line_center - page_center) < 50 and len(text) < 60
        
        # Check if this looks like a header/title
        is_header = (
            line['top'] < 150 and 
            len(text) < 50 and 
            not any(c.isdigit() for c in text[:4])
        )
        
        # Check if it's a date line (starts with "May", "June", etc.)
        is_date = bool(re.match(r'^(May|June|July|Aug|Sept|Oct|Nov|Dec|Jan|Feb|Mar|Apr)', text))
        
        analyzed.append({
            'text': text,
            'indent': indent_level,
            'centered': is_centered,
            'is_header': is_header,
            'is_date': is_date,
            'top': line['top']
        })
    
    return analyzed

def create_docx(analyzed_lines, output_path):
    """Create a Word document that mimics the original layout."""
    doc = Document()
    
    # Set up page margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # Use Courier New for typewriter look
    font_name = "Courier New"
    font_size = Pt(10)
    
    prev_top = 0
    
    for i, line_info in enumerate(analyzed_lines):
        text = line_info['text']
        
        # Skip empty lines
        if not text.strip():
            continue
        
        # Add blank paragraph for large vertical gaps
        if i > 0:
            gap = line_info['top'] - prev_top
            if gap > 50:  # Significant gap
                blank = doc.add_paragraph()
                blank.paragraph_format.space_after = Pt(0)
        
        # Create paragraph
        p = doc.add_paragraph()
        
        # Formatting based on analysis
        if line_info['centered']:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if line_info['indent'] > 0:
                p.paragraph_format.left_indent = Inches(line_info['indent'] * 0.25)
        
        # Tight line spacing for typewriter look
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        
        # Add the text
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = font_size
        
        # Bold for headers and date lines
        if line_info['is_header'] or line_info['is_date']:
            run.bold = False  # Keep consistent with typewriter look
        
        # Underline date lines (common in military docs)
        if line_info['is_date']:
            run.underline = True
        
        prev_top = line_info['top']
    
    doc.save(output_path)
    print(f"Created: {output_path}")

def main():
    if len(sys.argv) < 2:
        image_path = "cobia_5th_patrol_report/page_03_0741.jpg"
    else:
        image_path = sys.argv[1]
    
    output_path = "page_03_ocr.docx"
    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    
    print(f"Processing: {image_path}")
    
    # Get OCR data
    data, (img_width, img_height) = get_ocr_with_layout(image_path)
    
    # Reconstruct lines
    lines = reconstruct_lines(data, img_width)
    print(f"Found {len(lines)} lines of text")
    
    # Analyze layout
    analyzed = analyze_lines(lines, img_width)
    
    # Create Word document
    create_docx(analyzed, output_path)
    
    # Print extracted text for review
    print("\n--- Extracted Text (with layout markers) ---")
    for line in analyzed:
        indent_marker = ">" * line['indent'] if line['indent'] > 0 else ""
        center_marker = "[C]" if line['centered'] else ""
        date_marker = "[D]" if line['is_date'] else ""
        print(f"{indent_marker}{center_marker}{date_marker} {line['text']}")

if __name__ == '__main__':
    main()
